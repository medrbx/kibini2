import pandas as pd
import datetime as dt
from os.path import join
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import pandas as pd
import matplotlib.pyplot as plt

class Poldoc:
    def __init__(self,site='med',year=None):
        """
        Permet de charger les données pour génération des fiches d'analyse des collections.

        Args : 
            site (str) : "med" or "bus" or "col"
                by default "med" = médiathèque
            
            year (int) :
                si vide, valeur par défaut = année n-1
        """
        self.site = site
        self.year = year

        if year is None:
            self.year = dt.datetime.today().year - 1
        
        self.df = pd.read_excel(join("data",f"{lastyear}_synthese.xlsx"),sheet_name='mediatheque')
        
        if site == 'bus':
            self.df = pd.read_excel(join("data",f"{lastyear}_synthese.xlsx"),sheet_name='zebre')
        
        if site == 'col':
            self.df = pd.read_excel(join("data",f"{lastyear}_synthese.xlsx"),sheet_name='collectivites')
     

    def clean_dataset(self):
        """
        Une méthode pour nettoyer les valeurs du jeu de données et procéder à l'analyse
        """
        self.df = self.df[self.df['support']!='Périodique']
        
        self.df.loc[
            (self.df['collection_lib3']=='Livres précieux') &
            (self.df['collection_lib1']=='Jeunesse'),'collection_lib3'] = 'Livres précieux Petits enfants'
        
        self.df.loc[self.df['collection_lib4'].isna(),'collection_lib4'] = 'absence libellé'

        self.df.rename(columns={'nb_exemplaires_créés_dans_annee':'nb_acquisitions'},inplace=True)


    def get_tableau1(self,collection_level=None,collection_lib=None,support=True):
        """
        Permet de générer le tableau Volumétrie des collections


        Args :

            collection_level (str) : 
                Permet de choisir sur quel niveau on recherche les collection_lib

            collection_lib (list of string) :
                Accepte une liste de noms de collections

            support (Bool): par déufaut True
                Si False, enlève le type de support des valeurs aggrégées

                
        """
        
        if support==False:
            self.tab1 = self.df[self.df[collection_level].isin(collection_lib)].groupby(cols2keep_WithoutSupport)[vals2keep_tab1].sum()
        
        else:
            self.tab1 = self.df[self.df[collection_level].isin(collection_lib)].groupby(cols2keep)[vals2keep_tab1].sum()
    
            self.tab1['tx_acroissement (En %)'] = round((self.tab1['nb_acquisitions'] - self.tab1['nb_exemplaires_éliminés'])/self.tab1['nb_exemplaires_empruntables']*100,1)
            self.tab1['tx_renouvellement (En %)'] = round((self.tab1['nb_acquisitions'] / self.tab1['nb_exemplaires_empruntables'])*100,1)
        
        self.tab1 = self.tab1.T
        #return(self.tab1)

        #self.filename = self.tab1.columns[0][1]
        #self.filename = "_".join(self.tab1.columns.get_level_values(-2))

    def get_tableau2(self,collection_level,collection_lib,support=True):
        """
            Permet de générer le tableau Volumétrie des collections
        
                Args :
        
                    collection_level (str) : 
                        Permet de choisir sur quel niveau on recherche les collection_lib
        
                    collection_lib (list of string) :
                        Accepte une liste de noms de collections
        
                    support (Bool): par déufaut True
                        Si False, enlève le type de support des valeurs aggrégées            
        """
        
        if support == False:
            self.tab2 = self.df[self.df[collection_level].isin(collection_lib)].groupby(cols2keep_WithoutSupport)[vals2keep_tab2].sum()
        else:
            self.tab2 = self.df[self.df[collection_level].isin(collection_lib)].groupby(cols2keep)[vals2keep_tab2].sum()

        self.tab2[f'evolution_prets_{n2}-{nref} (En %)'] = round((self.tab2[f"nb_prets_{nref}"] - self.tab2[f'nb_prets_{n2}'])/self.tab2[f'nb_prets_{n2}']*100,1)
        self.tab2[f'evolution_prets_{n1}-{nref} (En %)'] = round((self.tab2[f"nb_prets_{nref}"] - self.tab2[f'nb_prets_{n1}'])/self.tab2[f'nb_prets_{n1}']*100,1)
        
        self.tab2[f'evolution_prets_emprunteurs_distincts_{n2}-{nref} (En %)'] = round((self.tab2[f'nb_prets_{nref}_emprunteurs_distincts'] - self.tab2[f'nb_prets_{n2}_emprunteurs_distincts']) / self.tab2[f'nb_prets_{n2}_emprunteurs_distincts']*100,1)
        self.tab2[f'evolution_prets_emprunteurs_distincts_{n1}-{nref} (En %)'] = round((self.tab2[f'nb_prets_{nref}_emprunteurs_distincts'] - self.tab2[f'nb_prets_{n1}_emprunteurs_distincts']) / self.tab2[f'nb_prets_{n1}_emprunteurs_distincts']*100,1)
        
        self.tab2['tx_rotation'] = (self.tab2[f'nb_prets_{nref}'] / self.tab2['nb_exemplaires_empruntables']).round(1)
        self.tab2['tx_sortie (En %)'] = (self.tab2['nb_exemplaires_en_pret'] / self.tab2['nb_exemplaires_empruntables']*100).round(1)
        self.tab2['tx_fonds_actifs (En %)'] = (self.tab2[f'nb_prets_{nref}_exemplaires_distincts'] / self.tab2['nb_exemplaires_empruntables']*100).round(1)
        self.tab2['Part de docs pas empruntés depuis 3 ans (En %)'] = (self.tab2['nb_exemplaires_empruntables_pas_empruntés_3_ans'] / self.tab2['nb_exemplaires_empruntables']*100).round(1)
        
        # Tri dans l'ordre souhaité
        self.tab2 = self.tab2[[f'nb_prets_{n2}',
                               f'nb_prets_{n1}',
                               f'nb_prets_{nref}',
                               f'evolution_prets_{n2}-{nref} (En %)',
                               f'evolution_prets_{n1}-{nref} (En %)',
                               f'nb_prets_{nref}_emprunteurs_distincts',
                               f'evolution_prets_emprunteurs_distincts_{n2}-{nref} (En %)',
                               f'evolution_prets_emprunteurs_distincts_{n1}-{nref} (En %)',
                               'tx_rotation',
                               'tx_fonds_actifs (En %)','Part de docs pas empruntés depuis 3 ans (En %)']]
        
        self.tab2 = self.tab2.T
        #return(tab2)

        
    def get_tableau3(self,secteur,collection_level,collection_lib):
        """
        Permet de générer le tableau Position dans l'Ensemble.

            Args :

                secteur (str) : Adultes or Jeunesse

                    Permet de choisir le Total secteur auquel les données de collections sont comparées.
                    Note : la variable est utilisée dans le nommage des fichiers en sortie pour les méthodes .save_tables et .get_FichesDanalyse

                collection_lib (list of string) :
                    Accepte une liste de noms de collections
    
                support (Bool): par déufaut True
                    Si False, enlève le type de support des valeurs aggrégées
            
        """
        
        self.secteur = secteur
        
        self.tab3 = self.df[self.df[collection_level].isin(collection_lib)]

        self.tab3 = self.tab3.pivot_table(index=collection_level,
                                                values=vals2keep_tab3,
                                                aggfunc=sum
                                                )
        self.tab3 = self.tab3.T

        # Ajout du total par ligne
        self.tab3.loc[self.tab3.index=='nb_acquisitions',f'Total_secteur_{secteur}'] = self.df[self.df['collection_lib1']==secteur]['nb_acquisitions'].sum()
        self.tab3.loc[self.tab3.index=='nb_exemplaires_empruntables',f'Total_secteur_{secteur}'] = self.df[self.df['collection_lib1']==secteur]['nb_exemplaires_empruntables'].sum()
        self.tab3.loc[self.tab3.index==f'nb_prets_{nref}',f'Total_secteur_{secteur}'] = self.df[self.df['collection_lib1']==secteur][f'nb_prets_{nref}'].sum()


        # Ajout de la part en %
        #tab3['Total_collections'] = tab3[collection_lib].sum(axis=1)
        self.tab3[f'Part collection {collection_lib} sur Total_{secteur} (En %)'] = round(self.tab3[collection_lib].sum(axis=1) / self.tab3[f'Total_secteur_{secteur}']*100,1)
     
        #return(tab3)


    def get_tableau4(self,keep_coll_lib=True):
        """
        Permet de sortir le tableau des budgets par collections.
        """
        if keep_coll_lib==False:
            self.tab4 = pd.DataFrame(index=['collection 1','collection 2','collection 3','etc...'],
                     columns=[f'budget_{self.year - 1}',
                              f'budget_{self.year}',
                              f'office_{self.year}',
                              f'commande_{self.year}',
                              f'suggestion_{self.year}',
                              f'evolution_budget_{self.year - 1}-{self.year}'
                             ])  

        else:
            self.tab4_coll_lib = self.tab2.columns.get_level_values(2).drop_duplicates()
            self.tab4 = pd.DataFrame(index=[lib for lib in self.tab4_coll_lib],
                                     columns=[f'budget_{self.year - 1}',
                                              f'budget_{self.year}',
                                              f'office_{self.year}',
                                              f'commande_{self.year}',
                                              f'suggestion_{self.year}',
                                              f'evolution_budget_{self.year - 1}-{self.year}'
                                             ]
                                    )

        


    
    def save_tables(self,
                    collection_filename='untitled',
                    in_docx=False):
        """
        Une méthode qui permet de transformer les dataframes en images et de les sauvegarder dans un répertoire
        """
        with pd.ExcelWriter(path=join('resultats','tableaux',f'tableaux_poldoc_{self.year}_{self.site}_{self.secteur}_{collection_filename}.xlsx')) as writer:
                self.tab1.to_excel(writer,sheet_name='VolumétrieDesCollections')
                self.tab2.to_excel(writer,sheet_name='UsageDesCollections')
                self.tab3.to_excel(writer,sheet_name='PositionDansEnsemble')
                self.tab4.to_excel(writer,sheet_name='Budgets (à remplir)')

        
        #On créer une liste regroupant nos 3 dataframes
        list_of_tables = [self.tab1,self.tab2,self.tab3]
        numb_of_tables = [1,2,3]
        df_filename = " & ".join(self.tab1.columns.get_level_values(-2)) 

        if in_docx:   
            document = Document()
            #sections = document.sections
            #sections.orientation = WD_ORIENT.LANDSCAPE
            #sections.page_width = 12
            #section.page_height = 5
            document.add_heading(f"Analyse des collections {self.year}\n{df_filename}",0)
            document.add_heading("0. Tableaux d'analyse")

            for table , numb in zip(list_of_tables,numb_of_tables):
                #df_display = table.reset_index(names='indicateur')
                df_display = table
                df_display.columns = [" - ".join([str(x) for x in col if x]) if isinstance(col, tuple) else col for col in df_display.columns]

                # Créer une figure vide
                fig, ax = plt.subplots(figsize=(8,2))
                ax.axis('off')  # cacher axes
                
                # Ajouter le DataFrame comme table
                tableau = ax.table(cellText=df_display,
                #colLabels=df_display.columns,
                colColours=["#e5c075"] * len(df_display.columns),
                #cellText=df_test.values,
                #colLabels=df_test.columns,
                loc='center',
                cellLoc='center')
                
                tableau.set_fontsize(50)
                tableau.scale(10,8)  # largeur/hauteur cellules
                #tableau.auto_set_font_size(True)
                
                
                # Sauvegarder en image
                plt.savefig(join("resultats","data",f"{self.year}_{self.site}_{df_filename}_tab_{numb}"),
                bbox_inches='tight',
                #dpi=1000
                )
                
                #plt.show()
                plt.close()
            
         
                
                document.add_picture(join("resultats","data",f"{self.year}_{self.site}_{df_filename}_tab_{numb}.png"),width=Inches(10.5))
            
            document.add_heading("1. Remarques générales",level=1)
            document.add_paragraph("Ecrivez quelque chose")
            document.add_heading("2. Raisons, hypothèses et contextualisation",level=1)
            document.add_paragraph("Caractéristiques des fonds, Faits marquants de l'année, (opération de désherbage, recotation, renouvellement / création de fonds, réaménagement, changement de responsable documentaire etc...)")
            document.add_paragraph("Ecrivez quelque chose")
            document.add_heading("3. Perpsectives points de vigilence",level=1)
            document.add_paragraph("Ecrivez quelque chose")


            for section in document.sections:
                # Passer en paysage
                section.orientation = WD_ORIENT.LANDSCAPE
                
                # Inverser largeur et hauteur
                new_width = Inches(12) #section.page_height
                new_height = Inches(8) #section.page_width
                section.page_width = new_width
                section.page_height = new_height
            
            document.save(join("resultats",
                               "fiches_poldoc",
                               f"{self.year}",f"fiche_analyse_{self.year}_{self.site}_{df_filename}.docx")
                         )
        
                

    def get_FichesDanalyse (self):
        """
        Permet de générer la liste de toutes les fiches d'analyse par secteur et par domaine documentaire.
        """

        liste_fiches_poldoc = ['med_Adultes_Arts&LoisirsCreatifs',
                               'med_Adultes_ArtsVivants',
                               'med_Adultes_Bd',
                               'med_Adultes_BdAuteurs',
                               'med_Adultes_BibsPro',
                               'med_Adultes_CineTV&SeriesTV',
                               'med_Adultes_CourtsRecits',
                               'med_Adultes_CreationGraphique',
                               'med_Adultes_FAL',
                               'med_Adultes_FilmsAutourDeRoubaix',
                               'med_Adultes_FondsLocalRegionalSonore',
                               'med_Adultes_FondsRegional',
                               'med_Adultes_FondsRegionalFilms',
                               'med_Adultes_FondsRegionalMusique',
                               'med_Adultes_HistoirePhiloReligion',
                               'med_Adultes_Humour',
                               'med_Adultes_Informatique',
                               'med_Adultes_LireAutrement',
                               'med_Adultes_LitteratureDeGenre',
                               'med_Adultes_LitteratureDivers',
                               'med_Adultes_Musique',
                               'med_Adultes_Parentalite',
                               'med_Adultes_Reussir&LanguesEtrangeres&FLE',
                               'med_Adultes_RomansAdos',
                               'med_Adultes_ViePratique&Loisirs',
                               'med_Adultes_ZéroDechet',
                            '   col_FondsPetiteEnfance',
                               'col_FondJeunesse',
                               'med_Jeunesse_Albums',
                               'med_Jeunesse_AlbumsDocumentaires',
                               'med_Jeunesse_AlbumsPE',
                               'med_Jeunesse_BD',
                               'med_Jeunesse_Cinema',
                               'med_Jeunesse_Musique',
                               'med_Jeunesse_Poesie&Théâtre',
                               'med_Jeunesse_Romans',
                               'bus_FondsAdultes',
                               'bus_FondsJeunesse'
                            ]
        
        for i in liste_fiches_poldoc:
            document = Document()
            document.add_heading(f'Analyse des collections {self.year}\n{i}',0)
            document.add_heading("1. Remarques générales",level=1)
            document.add_paragraph("Ecrivez quelque chose")
            document.add_heading("2. Raisons, hypothèses et contextualisation",level=1)
            document.add_paragraph("Caractéristiques des fonds, Faits marquants de l'année, (opération de désherbage, recotation, renouvellement / création de fonds, réaménagement, changement de responsable documentaire etc...)")
            document.add_paragraph("Ecrivez quelque chose")
            document.add_heading("3. Perpsectives points de vigilence",level=1)
            document.add_paragraph("Ecrivez quelque chose")
            document.save(join("resultats",f'{self.year}',f"fiche_analyse_{self.year}_{i}.docx"))